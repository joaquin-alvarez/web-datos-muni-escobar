#!/usr/bin/env python3
"""
Script to convert all .docx files to .md recursively in the data directory.
"""

import os
import sys
from pathlib import Path
from docx import Document
import re

def docx_to_markdown(docx_path):
    """Convert a .docx file to markdown format."""
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Simple conversion: treat as plain text for now
                # Could be enhanced to detect headings, lists, etc.
                markdown_content.append(text)
        
        # Extract tables
        for table in doc.tables:
            markdown_content.append("\n| " + " | ".join([cell.text.strip() for cell in table.rows[0].cells]) + " |")
            markdown_content.append("| " + " | ".join(["---" for _ in table.rows[0].cells]) + " |")
            
            for row in table.rows[1:]:
                markdown_content.append("| " + " | ".join([cell.text.strip() for cell in row.cells]) + " |")
            markdown_content.append("")
        
        return "\n".join(markdown_content)
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")
        return None

def convert_docx_files_recursive(data_dir):
    """Convert all .docx files to .md recursively in the given directory."""
    data_path = Path(data_dir)
    
    if not data_path.exists():
        print(f"Error: Directory {data_dir} does not exist")
        return
    
    converted_count = 0
    error_count = 0
    
    # Find all .docx files recursively
    docx_files = list(data_path.rglob("*.docx"))
    
    print(f"Found {len(docx_files)} .docx files to convert...")
    
    for docx_file in docx_files:
        # Create corresponding .md file path
        md_file = docx_file.with_suffix('.md')
        
        print(f"Converting: {docx_file.relative_to(data_path)} -> {md_file.relative_to(data_path)}")
        
        # Convert .docx to markdown
        markdown_content = docx_to_markdown(docx_file)
        
        if markdown_content is not None:
            try:
                # Write markdown content to file
                with open(md_file, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                converted_count += 1
                print(f"  ✓ Successfully converted")
            except Exception as e:
                print(f"  ✗ Error writing {md_file}: {e}")
                error_count += 1
        else:
            error_count += 1
    
    print(f"\nConversion complete!")
    print(f"Successfully converted: {converted_count} files")
    print(f"Errors: {error_count} files")

if __name__ == "__main__":
    data_directory = "/home/JA/development/web-datos-muni-escobar/data"
    
    if len(sys.argv) > 1:
        data_directory = sys.argv[1]
    
    print(f"Starting conversion of .docx files to .md in: {data_directory}")
    convert_docx_files_recursive(data_directory)
